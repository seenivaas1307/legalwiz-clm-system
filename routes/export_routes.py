# export_routes.py - PDF & DOCX Contract Export
"""
Export generated contracts as downloadable PDF or DOCX files.
Reuses the existing generate_contract() pipeline.

Endpoints:
- GET /{contract_id}/export/pdf  — Download as PDF
- GET /{contract_id}/export/docx — Download as DOCX
"""

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import StreamingResponse
from datetime import datetime
from io import BytesIO
import re

from config import DB_CONFIG
import psycopg2
from psycopg2.extras import RealDictCursor

router = APIRouter(prefix="/api/contracts", tags=["export"])


# ==================== HELPERS ====================

def get_db():
    return psycopg2.connect(**DB_CONFIG)


def _get_contract_data(contract_id: str) -> dict:
    """
    Get full generated contract data.
    Reuses the same pipeline as contract_generation_routes.generate_contract.
    """
    from contract_generation_routes import (
        get_active_clauses_with_text,
        get_parameter_values,
        get_parameter_names_map,
        replace_parameters,
    )

    conn = get_db()
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT id, title, contract_type, jurisdiction
                FROM contracts WHERE id = %s
            """, (contract_id,))
            contract = cur.fetchone()
            if not contract:
                raise HTTPException(status_code=404, detail="Contract not found")
    finally:
        conn.close()

    clauses = get_active_clauses_with_text(contract_id)
    if not clauses:
        raise HTTPException(
            status_code=400,
            detail="No active clauses found. Please complete Step 3 first."
        )

    param_values = get_parameter_values(contract_id)
    param_names = get_parameter_names_map(contract_id)

    processed = []
    all_missing = set()
    for clause in clauses:
        rendered, missing = replace_parameters(
            clause["raw_text"], param_values, param_names
        )
        processed.append({
            "clause_id": clause["clause_id"],
            "clause_type": clause["clause_type"],
            "variant": clause["variant"],
            "sequence": clause["sequence"],
            "rendered_text": rendered,
            "missing_parameters": missing,
        })
        all_missing.update(missing)

    return {
        "contract_id": contract_id,
        "title": contract.get("title", "Contract Agreement"),
        "contract_type": contract.get("contract_type", "General Agreement"),
        "jurisdiction": contract.get("jurisdiction", "India"),
        "generated_at": datetime.now(),
        "clauses": processed,
        "is_complete": len(all_missing) == 0,
        "missing_parameters": sorted(list(all_missing)),
    }


def _sanitize_filename(title: str) -> str:
    """Make a safe filename from contract title."""
    name = re.sub(r'[^\w\s-]', '', title).strip()
    name = re.sub(r'\s+', '_', name)
    return name[:80] or "contract"


# ==================== HTML BUILDER ====================

def _build_html(data: dict, watermark: str = None) -> str:
    """Build styled HTML for PDF conversion."""
    title = data["title"]
    parts = []
    parts.append("<!DOCTYPE html>")
    parts.append('<html lang="en"><head><meta charset="UTF-8">')
    parts.append(f"<title>{title}</title>")
    parts.append("""<style>
        @page {
            size: A4;
            margin: 2.5cm 2cm;
            @bottom-center { content: "Page " counter(page) " of " counter(pages); font-size: 10px; color: #888; }
        }
        body {
            font-family: 'Times New Roman', 'DejaVu Serif', Georgia, serif;
            font-size: 12pt;
            line-height: 1.6;
            color: #1a1a1a;
        }
        .header {
            text-align: center;
            border-bottom: 2px solid #000;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }
        .title { font-size: 22pt; font-weight: bold; margin-bottom: 8px; }
        .subtitle { font-size: 11pt; color: #555; margin-bottom: 4px; }
        .clause {
            margin-bottom: 24px;
            page-break-inside: avoid;
        }
        .clause-title {
            font-weight: bold;
            font-size: 13pt;
            margin-bottom: 8px;
            color: #111;
            border-bottom: 1px solid #ccc;
            padding-bottom: 4px;
        }
        .clause-text {
            text-align: justify;
            white-space: pre-wrap;
        }
        .missing-param {
            background-color: #FFEB3B;
            padding: 1px 4px;
            border-radius: 2px;
            font-weight: bold;
        }
        .watermark {
            position: fixed;
            top: 40%;
            left: 15%;
            font-size: 80pt;
            color: rgba(200, 200, 200, 0.3);
            transform: rotate(-45deg);
            z-index: -1;
            font-weight: bold;
        }
        .footer {
            text-align: center;
            margin-top: 40px;
            border-top: 2px solid #000;
            padding-top: 15px;
        }
        .signature-block {
            display: flex;
            justify-content: space-between;
            margin-top: 60px;
        }
        .sig-line {
            width: 40%;
            text-align: center;
        }
        .sig-line hr { border-top: 1px solid #000; margin-bottom: 5px; }
    </style>""")
    parts.append("</head><body>")

    # Watermark
    if watermark:
        parts.append(f'<div class="watermark">{watermark}</div>')

    # Header
    parts.append('<div class="header">')
    parts.append(f'<div class="title">{title}</div>')
    ctype = data["contract_type"].replace("_", " ").title()
    parts.append(f'<div class="subtitle">{ctype} | Jurisdiction: {data["jurisdiction"]}</div>')
    parts.append(f'<div class="subtitle">Generated: {data["generated_at"].strftime("%B %d, %Y")}</div>')
    parts.append("</div>")

    # Clauses
    for i, clause in enumerate(data["clauses"], 1):
        parts.append('<div class="clause">')
        clause_name = clause["clause_type"].replace("-", " ").replace("_", " ").title()
        parts.append(f'<div class="clause-title">{i}. {clause_name}</div>')

        text = clause["rendered_text"]
        # Highlight missing params
        for m in clause["missing_parameters"]:
            text = text.replace(m, f'<span class="missing-param">{m}</span>')
        # Convert newlines to <br>
        text = text.replace("\n\n", "</p><p>").replace("\n", "<br>")
        parts.append(f'<div class="clause-text"><p>{text}</p></div>')
        parts.append("</div>")

    # Signature block
    parts.append('<div class="signature-block">')
    parts.append('<div class="sig-line"><hr><p>Authorized Signatory (Party A)</p><p>Date: _______________</p></div>')
    parts.append('<div class="sig-line"><hr><p>Authorized Signatory (Party B)</p><p>Date: _______________</p></div>')
    parts.append("</div>")

    # Footer
    parts.append('<div class="footer">')
    parts.append("<p><strong>— END OF CONTRACT —</strong></p>")
    if not data["is_complete"]:
        parts.append(f'<p style="color: #c00;">⚠ {len(data["missing_parameters"])} parameter(s) still missing</p>')
    parts.append("</div>")

    parts.append("</body></html>")
    return "\n".join(parts)


# ==================== DOCX BUILDER ====================

def _build_docx(data: dict) -> BytesIO:
    """Build a DOCX document from contract data."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading(data["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Subtitle
    ctype = data["contract_type"].replace("_", " ").title()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{ctype} | Jurisdiction: {data['jurisdiction']}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub2.add_run(f"Generated: {data['generated_at'].strftime('%B %d, %Y')}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Separator
    doc.add_paragraph("─" * 60)

    # Clauses
    for i, clause in enumerate(data["clauses"], 1):
        clause_name = clause["clause_type"].replace("-", " ").replace("_", " ").title()
        heading = doc.add_heading(f"{i}. {clause_name}", level=2)
        for run in heading.runs:
            run.font.size = Pt(13)

        text = clause["rendered_text"]
        # Split into paragraphs on double newlines
        paragraphs = text.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Handle missing params — highlight in bold yellow-ish
            parts = re.split(r'(\{\{[A-Z_0-9]+\}\})', para_text)
            for part in parts:
                if re.match(r'\{\{[A-Z_0-9]+\}\}', part):
                    run = p.add_run(part)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)

    # Signature block
    doc.add_paragraph("")
    doc.add_paragraph("─" * 60)
    doc.add_paragraph("")

    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.columns[0].width = Inches(3)
    sig_table.columns[1].width = Inches(3)

    sig_table.cell(0, 0).text = "_________________________"
    sig_table.cell(0, 1).text = "_________________________"
    sig_table.cell(1, 0).text = "Authorized Signatory (Party A)"
    sig_table.cell(1, 1).text = "Authorized Signatory (Party B)"
    sig_table.cell(2, 0).text = "Date: _______________"
    sig_table.cell(2, 1).text = "Date: _______________"

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— END OF CONTRACT —")
    run.bold = True

    if not data["is_complete"]:
        warn = doc.add_paragraph()
        warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = warn.add_run(f"⚠ {len(data['missing_parameters'])} parameter(s) still missing")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ==================== ROUTES ====================

def _build_pdf(data: dict, watermark: str = None) -> BytesIO:
    """Build a PDF document using fpdf2 (pure Python, no C deps)."""
    from fpdf import FPDF

    class ContractPDF(FPDF):
        def __init__(self, title, watermark_text=None):
            super().__init__()
            self.contract_title = title
            self.watermark_text = watermark_text

        def header(self):
            if self.watermark_text and self.page_no() > 0:
                self.set_font("Helvetica", "B", 60)
                self.set_text_color(220, 220, 220)
                self.rotate(45, self.w / 2, self.h / 2)
                self.text(self.w / 4, self.h / 2, self.watermark_text)
                self.rotate(0)
                self.set_text_color(0, 0, 0)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", 0, 0, "C")

    pdf = ContractPDF(data["title"], watermark)
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, data["title"], 0, 1, "C")
    pdf.ln(4)

    # Subtitle
    ctype = data["contract_type"].replace("_", " ").title()
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"{ctype}  |  Jurisdiction: {data['jurisdiction']}", 0, 1, "C")
    pdf.cell(0, 6, f"Generated: {data['generated_at'].strftime('%B %d, %Y')}", 0, 1, "C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # Divider
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
    pdf.ln(10)

    # Clauses
    for i, clause in enumerate(data["clauses"], 1):
        clause_name = clause["clause_type"].replace("-", " ").replace("_", " ").title()

        # Clause heading
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, f"{i}. {clause_name}", 0, 1)
        pdf.ln(2)

        # Clause text
        pdf.set_font("Times", "", 11)
        text = clause["rendered_text"]
        # Handle encoding
        text = text.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 6, text)
        pdf.ln(4)

        # Light divider between clauses
        pdf.set_draw_color(200, 200, 200)
        pdf.set_line_width(0.2)
        pdf.line(30, pdf.get_y(), pdf.w - 30, pdf.get_y())
        pdf.ln(6)

    # Signature block
    pdf.ln(20)
    y_sig = pdf.get_y()
    pdf.set_font("Helvetica", "", 10)

    # Left signature
    pdf.set_xy(20, y_sig)
    pdf.cell(80, 6, "_________________________", 0, 1)
    pdf.set_x(20)
    pdf.cell(80, 6, "Authorized Signatory (Party A)", 0, 1)
    pdf.set_x(20)
    pdf.cell(80, 6, "Date: _______________", 0, 1)

    # Right signature
    pdf.set_xy(pdf.w - 100, y_sig)
    pdf.cell(80, 6, "_________________________", 0, 1)
    pdf.set_xy(pdf.w - 100, y_sig + 6)
    pdf.cell(80, 6, "Authorized Signatory (Party B)", 0, 1)
    pdf.set_xy(pdf.w - 100, y_sig + 12)
    pdf.cell(80, 6, "Date: _______________", 0, 1)

    # Footer
    pdf.ln(20)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
    pdf.ln(6)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "--- END OF CONTRACT ---", 0, 1, "C")

    if not data["is_complete"]:
        pdf.set_text_color(200, 0, 0)
        pdf.set_font("Helvetica", "I", 9)
        pdf.cell(0, 8, f"Warning: {len(data['missing_parameters'])} parameter(s) still missing", 0, 1, "C")

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


@router.get("/{contract_id}/export/pdf")
async def export_pdf(
    contract_id: str,
    watermark: str = Query(None, description="Optional watermark text (e.g. 'DRAFT')")
):
    """
    Export contract as a downloadable PDF file.

    Uses fpdf2 (pure Python) to generate a professional A4 PDF
    with page numbers, proper typography, and optional watermark.
    """
    data = _get_contract_data(contract_id)

    try:
        buf = _build_pdf(data, watermark=watermark)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="fpdf2 not installed. Run: pip install fpdf2"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"PDF generation failed: {str(e)}"
        )

    pdf_bytes = buf.getvalue()
    filename = _sanitize_filename(data["title"])

    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}.pdf"',
            "Content-Length": str(len(pdf_bytes)),
        }
    )


@router.get("/{contract_id}/export/docx")
async def export_docx(contract_id: str):
    """
    Export contract as a downloadable DOCX (Word) file.

    Uses python-docx to create a structured Word document with
    proper headings, formatting, and signature blocks.
    """
    data = _get_contract_data(contract_id)

    try:
        buf = _build_docx(data)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Run: pip install python-docx"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"DOCX generation failed: {str(e)}"
        )

    filename = _sanitize_filename(data["title"])
    content = buf.getvalue()

    return StreamingResponse(
        BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}.docx"',
            "Content-Length": str(len(content)),
        }
    )
